from pathlib import Path

import docx
from docx.table import Table as _DocxTable

from src.parsers.base import DocumentParser, ParsedDocument

_HEADING_LEVELS = {1: "# ", 2: "## ", 3: "### ", 4: "#### ", 5: "##### ", 6: "###### "}


def _format_table(table: _DocxTable) -> str:
    """Convert a docx table to a markdown-style string."""
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    # Pad rows to uniform width
    for r in rows:
        while len(r) < col_count:
            r.append("")

    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in range(col_count)) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join([header, sep, body])


class DocxParser(DocumentParser):
    def parse(self, path: Path) -> ParsedDocument:
        doc = docx.Document(str(path))
        parts: list[str] = []
        table_idx = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = docx.text.paragraph.Paragraph(element, doc)
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    try:
                        level = int(style.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    prefix = _HEADING_LEVELS.get(level, "# ")
                    parts.append(prefix + text)
                else:
                    parts.append(text)

            elif tag == "tbl":
                if table_idx < len(doc.tables):
                    md_table = _format_table(doc.tables[table_idx])
                    if md_table:
                        parts.append(md_table)
                    table_idx += 1

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"paragraphs": len(parts)},
            source_path=str(path),
            file_type="docx",
        )
