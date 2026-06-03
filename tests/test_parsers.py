import json
import tempfile
from pathlib import Path
import pytest
from src.parsers import parse_file


def test_text_parser():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("Hello world\nSecond line")
        f.flush()
        result = parse_file(Path(f.name))
        assert "Hello world" in result.content
        assert result.file_type == "text"


def test_markdown_parser():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write("# Title\n\nSome content")
        f.flush()
        result = parse_file(Path(f.name))
        assert "# Title" in result.content
        assert result.file_type == "markdown"


def test_unsupported_format():
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"data")
        f.flush()
        with pytest.raises(ValueError, match="Unsupported"):
            parse_file(Path(f.name))


def test_csv_parser_basic(tmp_path):
    csv_file = tmp_path / "test.csv"
    csv_file.write_text("name,age,city\nAlice,30,NYC\nBob,25,LA\n")
    result = parse_file(csv_file)
    assert "| name | age | city |" in result.content
    assert "| Alice | 30 | NYC |" in result.content
    assert "| Bob | 25 | LA |" in result.content
    assert result.file_type == "csv"


def test_csv_parser_tsv(tmp_path):
    tsv_file = tmp_path / "test.tsv"
    tsv_file.write_text("name\tage\tcity\nAlice\t30\tNYC\nBob\t25\tLA\n")
    result = parse_file(tsv_file)
    assert "| name | age | city |" in result.content
    assert "| Alice | 30 | NYC |" in result.content


def test_csv_parser_metadata(tmp_path):
    csv_file = tmp_path / "test.csv"
    csv_file.write_text("a,b\n1,2\n")
    result = parse_file(csv_file)
    assert result.metadata["delimiter"] == ","
    assert result.metadata["encoding"] == "utf-8"


def test_html_parser_basic(tmp_path):
    html_file = tmp_path / "test.html"
    html_file.write_text(
        "<html><head><title>My Page</title></head>"
        "<body><h1>Hello</h1><p>World</p></body></html>"
    )
    result = parse_file(html_file)
    assert "Hello" in result.content
    assert "World" in result.content
    assert "# Hello" in result.content
    assert result.file_type == "html"


def test_html_parser_strips_scripts(tmp_path):
    html_file = tmp_path / "test.html"
    html_file.write_text(
        "<html><body>"
        "<script>var x = 1;</script>"
        "<style>.cls { color: red; }</style>"
        "<nav>Navigation</nav>"
        "<footer>Footer stuff</footer>"
        "<p>Keep this</p>"
        "</body></html>"
    )
    result = parse_file(html_file)
    assert "var x" not in result.content
    assert "color: red" not in result.content
    assert "Navigation" not in result.content
    assert "Footer stuff" not in result.content
    assert "Keep this" in result.content


def test_html_parser_metadata(tmp_path):
    html_file = tmp_path / "test.html"
    html_file.write_text(
        "<html><head><title>Test Title</title></head><body>Content</body></html>"
    )
    result = parse_file(html_file)
    assert result.metadata["title"] == "Test Title"
    assert result.metadata["format"] == "html"


def test_docx_parser_basic(tmp_path):
    """Parse a DOCX with plain paragraphs and a table."""
    import docx

    docx_file = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Foo"
    table.cell(1, 1).text = "Bar"
    doc.save(str(docx_file))

    result = parse_file(docx_file)

    assert result.file_type == "docx"
    assert "First paragraph" in result.content
    assert "Second paragraph" in result.content
    # Table rendered as markdown
    assert "| Name | Value |" in result.content
    assert "| Foo | Bar |" in result.content
    assert "---" in result.content


def test_docx_parser_headings(tmp_path):
    """Verify heading hierarchy markers in output."""
    import docx

    docx_file = tmp_path / "headings.docx"
    doc = docx.Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text")
    doc.add_heading("Subsection", level=2)
    doc.add_heading("Detail", level=3)
    doc.save(str(docx_file))

    result = parse_file(docx_file)
    lines = result.content.split("\n\n")
    assert any(line.startswith("# Title") for line in lines)
    assert any(line.startswith("## Subsection") for line in lines)
    assert any(line.startswith("### Detail") for line in lines)
    assert "Body text" in result.content


def test_excel_parser_basic(tmp_path):
    """Parse an Excel file created programmatically."""
    import openpyxl

    xlsx_file = tmp_path / "test.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age", "City"])
    ws.append(["Alice", 30, "NYC"])
    ws.append(["Bob", 25, "LA"])
    wb.save(str(xlsx_file))
    wb.close()

    result = parse_file(xlsx_file)

    assert result.file_type == "excel"
    assert "## Sheet1" in result.content
    assert "Name | Age | City" in result.content
    assert "Alice | 30 | NYC" in result.content
    assert "Bob | 25 | LA" in result.content


def test_excel_parser_metadata(tmp_path):
    """Check that metadata includes sheet count and names."""
    import openpyxl

    xlsx_file = tmp_path / "meta.xlsx"
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Data"
    ws1.append(["x", "y"])
    ws2 = wb.create_sheet("Summary")
    ws2.append(["total", 42])
    wb.save(str(xlsx_file))
    wb.close()

    result = parse_file(xlsx_file)

    assert result.metadata["sheets"] == 2
    assert "Data" in result.metadata["sheet_names"]
    assert "Summary" in result.metadata["sheet_names"]


def test_markdown_parser_frontmatter(tmp_path: Path):
    content = (
        "---\ntitle: Hello\ntags:\n  - python\n  - rag\nauthor: test\n---\n\n"
        "# Heading\n\nBody text here.\n"
    )
    md_file = tmp_path / "with_meta.md"
    md_file.write_text(content, encoding="utf-8")

    result = parse_file(md_file)

    assert result.file_type == "markdown"
    assert result.metadata["frontmatter"] == {
        "title": "Hello",
        "tags": ["python", "rag"],
        "author": "test",
    }
    assert "---" not in result.content
    assert "# Heading" in result.content
    assert "Body text here." in result.content


def test_markdown_parser_no_frontmatter(tmp_path: Path):
    content = "# Title\n\nPlain content without front-matter.\n"
    md_file = tmp_path / "no_meta.md"
    md_file.write_text(content, encoding="utf-8")

    result = parse_file(md_file)

    assert result.file_type == "markdown"
    assert "frontmatter" not in result.metadata
    assert "# Title" in result.content
    assert "Plain content" in result.content


# --- JSON parser tests ---


def test_json_parser_object(tmp_path):
    """Parse a JSON object — flattened to key=value pairs."""
    json_file = tmp_path / "obj.json"
    data = {"name": "Alice", "age": 30, "address": {"city": "NYC", "zip": "10001"}}
    json_file.write_text(json.dumps(data), encoding="utf-8")

    result = parse_file(json_file)

    assert result.file_type == "json"
    assert "name=Alice" in result.content
    assert "age=30" in result.content
    assert "address.city=NYC" in result.content
    assert "address.zip=10001" in result.content


def test_json_parser_array(tmp_path):
    """Parse a JSON array of objects — rendered as markdown table."""
    json_file = tmp_path / "arr.json"
    data = [
        {"name": "Alice", "age": 30},
        {"name": "Bob", "age": 25},
    ]
    json_file.write_text(json.dumps(data), encoding="utf-8")

    result = parse_file(json_file)

    assert result.file_type == "json"
    assert "| name | age |" in result.content
    assert "| Alice | 30 |" in result.content
    assert "| Bob | 25 |" in result.content
    assert "---" in result.content


def test_json_parser_metadata(tmp_path):
    """Check metadata reports number of keys for objects."""
    json_file = tmp_path / "meta.json"
    data = {"a": 1, "b": 2, "c": 3}
    json_file.write_text(json.dumps(data), encoding="utf-8")

    result = parse_file(json_file)

    assert result.file_type == "json"
    assert result.metadata["keys"] == 3


# ---------------------------------------------------------------------------
# PDF parser tests
# ---------------------------------------------------------------------------


def _make_simple_pdf(path: Path, text: str = "Hello PDF World") -> Path:
    """Create a minimal valid PDF file containing *text*."""
    content_stream = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET"
    stream_bytes = content_stream.encode("latin-1")
    stream_length = len(stream_bytes)

    objects: list[bytes] = []

    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objects.append(
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    objects.append(
        f"4 0 obj\n<< /Length {stream_length} >>\nstream\n".encode("latin-1")
        + stream_bytes
        + b"\nendstream\nendobj\n"
    )
    objects.append(
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    )

    pdf_bytes = b"%PDF-1.4\n"
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(pdf_bytes))
        pdf_bytes += obj

    xref_offset = len(pdf_bytes)
    pdf_bytes += b"xref\n"
    pdf_bytes += f"0 {len(objects) + 1}\n".encode()
    pdf_bytes += b"0000000000 65535 f \n"
    for off in offsets:
        pdf_bytes += f"{off:010d} 00000 n \n".encode()
    pdf_bytes += b"trailer\n"
    pdf_bytes += f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    pdf_bytes += b"startxref\n"
    pdf_bytes += f"{xref_offset}\n".encode()
    pdf_bytes += b"%%EOF\n"

    path.write_bytes(pdf_bytes)
    return path


def test_pdf_parser_basic(tmp_path):
    pdf_path = _make_simple_pdf(tmp_path / "test.pdf")
    result = parse_file(pdf_path)
    assert "Hello" in result.content or "PDF" in result.content
    assert result.file_type == "pdf"
    assert result.source_path == str(pdf_path)


def test_pdf_parser_metadata(tmp_path):
    pdf_path = _make_simple_pdf(tmp_path / "test.pdf")
    result = parse_file(pdf_path)
    assert result.metadata["pages"] == 1
    assert "tables_found" in result.metadata
